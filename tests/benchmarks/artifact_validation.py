from __future__ import annotations

import re
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import asdict, dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Literal, Sequence
from xml.etree import ElementTree

ArtifactType = Literal["ppt", "docx", "html-game"]

OFFICE_REQUIRED_ENTRIES = {
    "ppt": {"[Content_Types].xml", "ppt/presentation.xml"},
    "docx": {"[Content_Types].xml", "word/document.xml"},
}


@dataclass
class ValidationResult:
    artifact_type: ArtifactType
    technical_pass: bool
    quality_pass: bool
    quality_score: float
    content_recall_pct: float
    checks: dict[str, bool] = field(default_factory=dict)
    metrics: dict[str, int | float | str | bool] = field(default_factory=dict)
    matched_term_groups: list[bool] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", "", text).casefold()


def _term_group_hits(text: str, groups: Sequence[Sequence[str]]) -> list[bool]:
    normalized = _normalize(text)
    return [any(_normalize(term) in normalized for term in group) for group in groups]


def _ratio(hits: Sequence[bool]) -> float:
    return sum(hits) / len(hits) if hits else 1.0


def _parse_xml_entries(archive: zipfile.ZipFile) -> tuple[bool, list[str]]:
    errors: list[str] = []
    for name in archive.namelist():
        if not name.lower().endswith((".xml", ".rels")):
            continue
        try:
            ElementTree.fromstring(archive.read(name))
        except Exception as exc:
            errors.append(f"Malformed XML {name}: {exc.__class__.__name__}")
    return not errors, errors


def _run_office_schema_validator(path: Path, backend_root: Path) -> tuple[bool, str]:
    script = backend_root / "skills" / "docx" / "scripts" / "office" / "validate.py"
    try:
        completed = subprocess.run(
            [sys.executable, str(script), str(path)],
            cwd=script.parent,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        return False, exc.__class__.__name__
    output = "\n".join(part.strip() for part in (completed.stdout, completed.stderr) if part.strip())
    return completed.returncode == 0, output[-1000:]


def validate_pptx(
    path: Path,
    *,
    required_term_groups: Sequence[Sequence[str]],
    backend_root: Path,
) -> ValidationResult:
    checks: dict[str, bool] = {"extension": path.suffix.lower() == ".pptx", "non_empty": path.stat().st_size > 0}
    errors: list[str] = []
    text_parts: list[str] = []
    slide_count = 0
    non_empty_slides = 0
    aspect_ratio = 0.0
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            checks["zip_integrity"] = archive.testzip() is None
            checks["required_entries"] = OFFICE_REQUIRED_ENTRIES["ppt"].issubset(names)
            checks["xml_parse"] , xml_errors = _parse_xml_entries(archive)
            errors.extend(xml_errors)
            slide_names = sorted(
                name for name in names if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            )
            slide_count = len(slide_names)
            for name in slide_names:
                root = ElementTree.fromstring(archive.read(name))
                slide_text = " ".join((element.text or "") for element in root.iter() if element.tag.endswith("}t"))
                if slide_text.strip():
                    non_empty_slides += 1
                    text_parts.append(slide_text)
            if "ppt/presentation.xml" in names:
                presentation = ElementTree.fromstring(archive.read("ppt/presentation.xml"))
                size = next((element for element in presentation.iter() if element.tag.endswith("}sldSz")), None)
                if size is not None:
                    width = int(size.attrib.get("cx", "0") or 0)
                    height = int(size.attrib.get("cy", "0") or 0)
                    aspect_ratio = width / height if height else 0.0
    except (OSError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
        checks["zip_integrity"] = False
        checks["required_entries"] = False
        checks["xml_parse"] = False
        errors.append(exc.__class__.__name__)

    checks["slide_count_8_to_12"] = 8 <= slide_count <= 12
    checks["non_empty_slides"] = slide_count > 0 and non_empty_slides / slide_count >= 0.9
    checks["aspect_ratio_16_9"] = abs(aspect_ratio - (16 / 9)) <= 0.05
    schema_ok, schema_output = _run_office_schema_validator(path, backend_root)
    checks["office_schema"] = schema_ok
    if not schema_ok:
        errors.append(f"Office schema validation failed: {schema_output}")

    term_hits = _term_group_hits("\n".join(text_parts), required_term_groups)
    content_ratio = _ratio(term_hits)
    structure_checks = [checks["slide_count_8_to_12"], checks["non_empty_slides"]]
    usability_checks = [checks["aspect_ratio_16_9"], checks["office_schema"]]
    quality_score = round(35 * content_ratio + 30 * _ratio(structure_checks) + 25 * _ratio(usability_checks) + 10, 2)
    technical_keys = ("extension", "non_empty", "zip_integrity", "required_entries", "xml_parse", "office_schema")
    technical_pass = all(checks.get(key, False) for key in technical_keys)
    return ValidationResult(
        artifact_type="ppt",
        technical_pass=technical_pass,
        quality_pass=technical_pass and content_ratio >= 0.75 and quality_score >= 80,
        quality_score=quality_score,
        content_recall_pct=round(content_ratio * 100, 2),
        checks=checks,
        metrics={
            "size_bytes": path.stat().st_size,
            "slide_count": slide_count,
            "non_empty_slide_count": non_empty_slides,
            "aspect_ratio": round(aspect_ratio, 4),
            "extracted_text_chars": len("".join(text_parts)),
        },
        matched_term_groups=term_hits,
        errors=errors,
    )


def validate_docx(
    path: Path,
    *,
    required_term_groups: Sequence[Sequence[str]],
    backend_root: Path,
) -> ValidationResult:
    checks: dict[str, bool] = {"extension": path.suffix.lower() == ".docx", "non_empty": path.stat().st_size > 0}
    errors: list[str] = []
    paragraphs: list[str] = []
    table_paragraphs: list[str] = []
    heading_count = 0
    table_count = 0
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            checks["zip_integrity"] = archive.testzip() is None
            checks["required_entries"] = OFFICE_REQUIRED_ENTRIES["docx"].issubset(names)
            checks["xml_parse"], xml_errors = _parse_xml_entries(archive)
            errors.extend(xml_errors)
        from docx import Document

        document = Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_paragraphs = [
            paragraph.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        ]
        heading_count = sum(
            1 for paragraph in document.paragraphs if str(getattr(paragraph.style, "name", "")).lower().startswith("heading")
        )
        table_count = len(document.tables)
        checks["python_docx_parse"] = True
    except Exception as exc:
        checks["zip_integrity"] = False
        checks["required_entries"] = False
        checks["xml_parse"] = False
        checks["python_docx_parse"] = False
        errors.append(exc.__class__.__name__)

    full_text = "\n".join([*paragraphs, *table_paragraphs])
    section_groups = [
        ("教学目标", "学习目标"),
        ("教学重点", "重点"),
        ("教学过程", "教学流程", "课堂流程"),
        ("评价", "学习证据"),
        ("作业", "课后任务", "延伸任务"),
    ]
    section_hits = _term_group_hits(full_text, section_groups)
    checks["required_sections"] = _ratio(section_hits) >= 0.8
    checks["structured_content"] = heading_count >= 3 and len(paragraphs) >= 10
    checks["table_or_rich_structure"] = table_count >= 1 or heading_count >= 5
    schema_ok, schema_output = _run_office_schema_validator(path, backend_root)
    checks["office_schema"] = schema_ok
    if not schema_ok:
        errors.append(f"Office schema validation failed: {schema_output}")

    term_hits = _term_group_hits(full_text, required_term_groups)
    content_ratio = _ratio(term_hits)
    quality_score = round(
        35 * content_ratio
        + 30 * _ratio(section_hits)
        + 25 * _ratio([checks["structured_content"], checks["table_or_rich_structure"]])
        + 10,
        2,
    )
    technical_keys = (
        "extension",
        "non_empty",
        "zip_integrity",
        "required_entries",
        "xml_parse",
        "python_docx_parse",
        "office_schema",
    )
    technical_pass = all(checks.get(key, False) for key in technical_keys)
    return ValidationResult(
        artifact_type="docx",
        technical_pass=technical_pass,
        quality_pass=technical_pass and content_ratio >= 0.75 and quality_score >= 80,
        quality_score=quality_score,
        content_recall_pct=round(content_ratio * 100, 2),
        checks=checks,
        metrics={
            "size_bytes": path.stat().st_size,
            "paragraph_count": len(paragraphs),
            "table_paragraph_count": len(table_paragraphs),
            "heading_count": heading_count,
            "table_count": table_count,
            "extracted_text_chars": len(full_text),
        },
        matched_term_groups=term_hits,
        errors=errors,
    )


class _HTMLInspector(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.tags: list[str] = []
        self.attrs: list[dict[str, str | None]] = []
        self.scripts: list[str] = []
        self._script_parts: list[str] | None = None
        self.text_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self.tags.append(tag)
        self.attrs.append(dict(attrs))
        if tag == "script":
            self._script_parts = []

    def handle_endtag(self, tag: str) -> None:
        if tag == "script" and self._script_parts is not None:
            self.scripts.append("".join(self._script_parts))
            self._script_parts = None

    def handle_data(self, data: str) -> None:
        if self._script_parts is not None:
            self._script_parts.append(data)
        else:
            self.text_parts.append(data)


def _javascript_syntax_ok(scripts: Sequence[str]) -> tuple[bool, str]:
    combined = "\n".join(script for script in scripts if script.strip())
    if not combined:
        return False, "No inline JavaScript"
    with tempfile.NamedTemporaryFile(mode="w", suffix=".js", encoding="utf-8", delete=False) as handle:
        handle.write(combined)
        temp_path = Path(handle.name)
    try:
        completed = subprocess.run(
            ["node", "--check", str(temp_path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
            check=False,
        )
        return completed.returncode == 0, completed.stderr[-1000:]
    except (OSError, subprocess.TimeoutExpired) as exc:
        return False, exc.__class__.__name__
    finally:
        temp_path.unlink(missing_ok=True)


def validate_html(path: Path, *, required_term_groups: Sequence[Sequence[str]]) -> ValidationResult:
    checks: dict[str, bool] = {"extension": path.suffix.lower() == ".html", "non_empty": path.stat().st_size > 0}
    errors: list[str] = []
    try:
        html = path.read_text(encoding="utf-8")
        inspector = _HTMLInspector()
        inspector.feed(html)
        checks["html_structure"] = all(tag in inspector.tags for tag in ("html", "head", "body"))
    except Exception as exc:
        html = ""
        inspector = _HTMLInspector()
        checks["html_structure"] = False
        errors.append(exc.__class__.__name__)

    controls = sum(tag in {"button", "input", "select", "textarea"} for tag in inspector.tags)
    script_text = "\n".join(inspector.scripts)
    syntax_ok, syntax_error = _javascript_syntax_ok(inspector.scripts)
    checks["javascript_syntax"] = syntax_ok
    if not syntax_ok:
        errors.append(f"JavaScript validation failed: {syntax_error}")
    checks["interactive_controls"] = controls >= 1
    checks["interaction_logic"] = any(
        marker in script_text for marker in ("addEventListener", "onclick", "oninput", "querySelector", "getElementById")
    )
    visible_text = " ".join(inspector.text_parts)
    checks["instant_feedback"] = any(term in visible_text + script_text for term in ("反馈", "正确", "错误", "再试", "提示"))
    checks["responsive_meta"] = "viewport" in html.casefold()
    checks["accessible_controls"] = any(
        attrs.get("aria-label") or attrs.get("title") or attrs.get("id") for attrs in inspector.attrs
    )
    checks["self_contained"] = not bool(re.search(r"<(?:script|link)[^>]+(?:src|href)=[\"']https?://", html, re.I))
    term_hits = _term_group_hits(visible_text, required_term_groups)
    content_ratio = _ratio(term_hits)
    pedagogical_hits = _term_group_hits(
        visible_text,
        [("学习目标", "教学目标"), ("任务", "活动", "实验"), ("总结", "要点"), ("反馈", "提示")],
    )
    usability = [
        checks["interactive_controls"],
        checks["interaction_logic"],
        checks["instant_feedback"],
        checks["self_contained"],
    ]
    accessibility = [checks["responsive_meta"], checks["accessible_controls"]]
    quality_score = round(
        35 * content_ratio + 30 * _ratio(pedagogical_hits) + 25 * _ratio(usability) + 10 * _ratio(accessibility),
        2,
    )
    technical_keys = ("extension", "non_empty", "html_structure", "javascript_syntax")
    technical_pass = all(checks.get(key, False) for key in technical_keys)
    return ValidationResult(
        artifact_type="html-game",
        technical_pass=technical_pass,
        quality_pass=technical_pass and content_ratio >= 0.75 and quality_score >= 80,
        quality_score=quality_score,
        content_recall_pct=round(content_ratio * 100, 2),
        checks=checks,
        metrics={
            "size_bytes": path.stat().st_size,
            "html_tag_count": len(inspector.tags),
            "script_count": len(inspector.scripts),
            "interactive_control_count": controls,
            "extracted_text_chars": len(visible_text),
        },
        matched_term_groups=term_hits,
        errors=errors,
    )


def validate_artifact(
    path: Path,
    *,
    artifact_type: ArtifactType,
    required_term_groups: Sequence[Sequence[str]],
    backend_root: Path,
) -> ValidationResult:
    if artifact_type == "ppt":
        return validate_pptx(path, required_term_groups=required_term_groups, backend_root=backend_root)
    if artifact_type == "docx":
        return validate_docx(path, required_term_groups=required_term_groups, backend_root=backend_root)
    return validate_html(path, required_term_groups=required_term_groups)
