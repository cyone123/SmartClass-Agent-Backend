from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from tests.benchmarks.artifact_validation import (
    _run_office_schema_validator,
    validate_artifact,
    validate_html,
)


def test_html_validator_accepts_self_contained_interaction(tmp_path: Path) -> None:
    artifact = tmp_path / "activity.html"
    artifact.write_text(
        """<!doctype html><html><head><meta name="viewport" content="width=device-width"></head>
        <body><h1>冒泡排序</h1><p>学习目标</p><p>比较相邻元素并交换</p>
        <button id="next" aria-label="下一步">下一步</button><div id="feedback">即时反馈</div>
        <script>document.getElementById('next').addEventListener('click', () => {
        document.getElementById('feedback').textContent = '正确'; });</script></body></html>""",
        encoding="utf-8",
    )

    result = validate_html(
        artifact,
        required_term_groups=[["冒泡排序"], ["相邻元素"], ["比较"], ["交换"]],
    )

    assert result.technical_pass is True
    assert result.quality_pass is True
    assert result.content_recall_pct == 100


def test_html_validator_rejects_invalid_javascript(tmp_path: Path) -> None:
    artifact = tmp_path / "broken.html"
    artifact.write_text(
        "<html><head></head><body><button>go</button><script>function broken( {</script></body></html>",
        encoding="utf-8",
    )

    result = validate_html(artifact, required_term_groups=[])

    assert result.technical_pass is False
    assert result.checks["javascript_syntax"] is False


def test_docx_validator_opens_real_document(tmp_path: Path, monkeypatch) -> None:
    artifact = tmp_path / "lesson.docx"
    document = Document()
    for title in ("教学目标", "教学重点", "教学过程", "形成性评价", "课后作业"):
        document.add_heading(title, level=1)
        document.add_paragraph("勾股定理用于直角三角形和校园旗杆测高。")
    document.save(artifact)

    monkeypatch.setattr(
        "tests.benchmarks.artifact_validation._run_office_schema_validator",
        lambda path, backend_root: (True, "ok"),
    )
    result = validate_artifact(
        artifact,
        artifact_type="docx",
        required_term_groups=[["勾股定理"], ["直角三角形"], ["旗杆"]],
        backend_root=tmp_path,
    )

    assert result.technical_pass is True
    assert result.content_recall_pct == 100
    assert result.metrics["heading_count"] == 5


def test_docx_validator_includes_table_cells_in_content_recall(tmp_path: Path, monkeypatch) -> None:
    artifact = tmp_path / "table-lesson.docx"
    document = Document()
    for title in ("教学目标", "教学重点", "教学过程", "形成性评价", "课后作业"):
        document.add_heading(title, level=1)
        document.add_paragraph("课堂活动说明。")
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "勾股定理"
    table.rows[0].cells[1].text = "直角三角形"
    table.rows[0].cells[2].text = "旗杆测高"
    document.save(artifact)

    monkeypatch.setattr(
        "tests.benchmarks.artifact_validation._run_office_schema_validator",
        lambda path, backend_root: (True, "ok"),
    )
    result = validate_artifact(
        artifact,
        artifact_type="docx",
        required_term_groups=[["勾股定理"], ["直角三角形"], ["旗杆"]],
        backend_root=tmp_path,
    )

    assert result.content_recall_pct == 100
    assert result.metrics["table_paragraph_count"] == 3


def test_office_schema_validator_does_not_decode_chinese_ooxml_as_gbk(tmp_path: Path) -> None:
    artifact = tmp_path / "chinese.docx"
    document = Document()
    document.add_heading("教学目标", level=1)
    document.add_paragraph("学生能够说明勾股定理，并完成校园旗杆测高任务。")
    document.save(artifact)

    backend_root = Path(__file__).resolve().parents[2]
    _, output = _run_office_schema_validator(artifact, backend_root)

    assert "codec can't decode" not in output


def test_pptx_validator_rejects_incomplete_package(tmp_path: Path) -> None:
    artifact = tmp_path / "broken.pptx"
    with zipfile.ZipFile(artifact, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types></Types>")

    result = validate_artifact(
        artifact,
        artifact_type="ppt",
        required_term_groups=[],
        backend_root=tmp_path,
    )

    assert result.technical_pass is False
    assert result.checks["required_entries"] is False
