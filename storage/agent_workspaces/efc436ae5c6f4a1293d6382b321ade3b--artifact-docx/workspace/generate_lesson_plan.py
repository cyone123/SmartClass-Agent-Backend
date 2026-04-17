#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Get output directory from environment
output_dir = os.environ.get('AGENT_OUTPUT_DIR', '.')
output_path = os.path.join(output_dir, '牛顿第一定律_教学设计.docx')

# Create document
doc = Document()

# Set document default font to Chinese
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# Helper function to set Chinese font
def set_chinese_font(run, font_name='SimSun', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Helper function to add heading
def add_heading_cn(doc, text, level=1, font_name='SimHei'):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, font_name, 16 if level == 1 else 14, bold=True)
    return heading

# Helper function to add paragraph with Chinese font
def add_paragraph_cn(doc, text, font_name='SimSun', size=12, bold=False, alignment=None, first_line_indent=None):
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    if first_line_indent:
        para.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = para.add_run(text)
    set_chinese_font(run, font_name, size, bold)
    return para

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('教学设计方案')
set_chinese_font(run, 'SimHei', 22, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('牛顿第一定律')
set_chinese_font(run, 'SimHei', 18, bold=True)

# 一、基本信息
add_heading_cn(doc, '一、基本信息', 1)

# Basic info table
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'

# Row 1
cells = table.rows[0].cells
cells[0].text = '课程名称'
cells[1].text = '牛顿第一定律'
cells[2].text = '适用年级'
cells[3].text = '初三'

# Row 2
cells = table.rows[1].cells
cells[0].text = '学科'
cells[1].text = '物理'
cells[2].text = '课时安排'
cells[3].text = '1课时（40分钟）'

# Row 3
cells = table.rows[2].cells
cells[0].text = '教学方法'
cells[1].text = '实验探究法'
cells[2].text = '教学用具'
cells[3].text = '斜面、小车、毛巾、棉布、木板、刻度尺、多媒体课件'

# Format table cells
for row in table.rows:
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                if i % 2 == 0:  # Header cells
                    set_chinese_font(run, 'SimHei', 11, bold=True)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    set_chinese_font(run, 'SimSun', 11)

# 二、教学目标
add_heading_cn(doc, '二、教学目标', 1)

add_heading_cn(doc, '1. 知识与技能', 2)
para = doc.add_paragraph(style='List Number')
run = para.add_run('理解牛顿第一定律的内容，能够复述定律并解释"一切"、"没有受到外力"、"总保持"等关键词的含义。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Number')
run = para.add_run('理解惯性是物体的固有属性，知道质量是惯性大小的量度。')
set_chinese_font(run, 'SimSun', 12)

add_heading_cn(doc, '2. 过程与方法', 2)
para = doc.add_paragraph(style='List Bullet')
run = para.add_run('通过"斜面小车"实验，体验"控制变量法"和"科学推理法"（理想实验法）在物理探究中的应用。')
set_chinese_font(run, 'SimSun', 12)

add_heading_cn(doc, '3. 情感态度与价值观', 2)
para = doc.add_paragraph(style='List Bullet')
run = para.add_run('通过对亚里士多德与伽利略观点的辨析，培养敢于质疑、实事求是的科学态度。')
set_chinese_font(run, 'SimSun', 12)

# 三、教学重难点
add_heading_cn(doc, '三、教学重难点', 1)

para = doc.add_paragraph()
run1 = para.add_run('【重点】')
set_chinese_font(run1, 'SimHei', 12, bold=True)
run2 = para.add_run('牛顿第一定律的建立过程及内容理解。')
set_chinese_font(run2, 'SimSun', 12)

para = doc.add_paragraph()
run1 = para.add_run('【难点】')
set_chinese_font(run1, 'SimHei', 12, bold=True)
run2 = para.add_run('对惯性概念的理解与辨析（特别是"物体运动状态改变与惯性无关"这一认知误区的突破）。')
set_chinese_font(run2, 'SimSun', 12)

# 四、教学过程设计
add_heading_cn(doc, '四、教学过程设计', 1)

# 环节一
add_heading_cn(doc, '环节一：情境导入与认知冲突（5分钟）', 2)

para = doc.add_paragraph()
run = para.add_run('【活动设计】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Number')
run = para.add_run('教师演示：用力推讲台上的粉笔盒，它运动；停止推力，它静止。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Number')
run = para.add_run('提问："物体的运动需要力来维持吗？"')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Number')
run = para.add_run('引导学生联想生活经验（如推车、踢球），大部分学生会直觉认同亚里士多德的观点："力是维持物体运动的原因"。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run1 = para.add_run('【教师引导】')
set_chinese_font(run1, 'SimHei', 12, bold=True)
run2 = para.add_run('展示亚里士多德与伽利略的观点对比，指出直觉往往是不可靠的，引入"实验探究"来验证。')
set_chinese_font(run2, 'SimSun', 12)

# 环节二
add_heading_cn(doc, '环节二：实验探究——牛顿第一定律的建立（15分钟）', 2)

para = doc.add_paragraph()
run = para.add_run('【实验操作】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('使用斜面小车装置。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('控制变量：让小车从斜面同一高度滑下（控制初速度相同）。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('改变条件：分别在毛巾表面、棉布表面、木板表面滑行。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【观察记录】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('引导学生观察小车在不同表面滑行的距离，并记录在表格中。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('数据分析：阻力越小，滑行距离越远，速度减小得越慢。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【科学推理】（关键步骤）')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('提问："如果木板绝对光滑（没有阻力），小车将运动多远？"')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('推理结论：小车将永远运动下去，做匀速直线运动。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【总结定律】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph()
run = para.add_run('综合伽利略、笛卡尔、牛顿的研究，引出牛顿第一定律：')
set_chinese_font(run, 'SimSun', 12)

# Law highlight box
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('一切物体在没有受到外力作用时，总保持静止状态或匀速直线运动状态。')
set_chinese_font(run, 'SimHei', 12, bold=True)

# 环节三
add_heading_cn(doc, '环节三：难点突破——惯性的理解与应用（15分钟）', 2)

para = doc.add_paragraph()
run = para.add_run('【概念界定】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('定律引申：物体保持运动状态不变的性质叫做惯性。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('强调：惯性是物体的固有属性，一切物体在任何情况下（无论运动、静止、受力、不受力）都有惯性。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【互动辨析】（难点攻克）')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('提问："高速行驶的汽车刹车时有惯性，静止时有没有惯性？"')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('纠错：纠正"物体运动越快惯性越大"的错误观点（惯性只与质量有关）。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【实际问题解决】（达成教学目标）')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Bullet')
run1 = para.add_run('案例1：')
set_chinese_font(run1, 'SimSun', 12, bold=True)
run2 = para.add_run('为什么锤头松了，把锤柄在地上撞几下就紧了？（解释：锤柄受力停止，锤头由于惯性继续向下运动，套紧。）')
set_chinese_font(run2, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run1 = para.add_run('案例2：')
set_chinese_font(run1, 'SimSun', 12, bold=True)
run2 = para.add_run('为什么跳远运动员起跳前要助跑？（解释：利用惯性保持向前的速度。）')
set_chinese_font(run2, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run1 = para.add_run('案例3：')
set_chinese_font(run1, 'SimSun', 12, bold=True)
run2 = para.add_run('交通安全中的系安全带是为了防止什么？（解释：防止急刹车时人由于惯性向前倾倒造成伤害。）')
set_chinese_font(run2, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【方法总结】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph()
run = para.add_run('引导学生总结解释惯性现象的三步法：')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('确认原状态；')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('指出受力物体运动状态改变；')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run('另一物体由于惯性保持原状态。')
set_chinese_font(run, 'SimSun', 12)

# 环节四
add_heading_cn(doc, '环节四：课堂小结与评估（5分钟）', 2)

para = doc.add_paragraph()
run = para.add_run('【知识梳理】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph()
run = para.add_run('回顾力的作用（改变运动状态而非维持运动）、牛顿第一定律内容、惯性概念。')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph()
run = para.add_run('【课堂提问评估】')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = doc.add_paragraph(style='List Number')
run = para.add_run('"牛顿第一定律是通过实验直接得出的吗？"（考查科学方法）')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Number')
run = para.add_run('"把衣服抖一下，灰尘为什么掉落？"（考查实际应用能力）')
set_chinese_font(run, 'SimSun', 12)

para = doc.add_paragraph(style='List Number')
run = para.add_run('"如果物体受力平衡，它会保持什么状态？"（为下一节二力平衡做铺垫）')
set_chinese_font(run, 'SimSun', 12)

# 五、板书设计
add_heading_cn(doc, '五、板书设计', 1)

# Blackboard design table
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]

# Add content to the cell
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('牛顿第一定律')
set_chinese_font(run, 'SimHei', 14, bold=True)

para = cell.add_paragraph()
run = para.add_run('1. 历史观点：')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = cell.add_paragraph()
run = para.add_run('    • 亚里士多德：力是维持运动的原因（错）')
set_chinese_font(run, 'SimSun', 12)

para = cell.add_paragraph()
run = para.add_run('    • 伽利略：力是改变运动状态的原因（对）')
set_chinese_font(run, 'SimSun', 12)

para = cell.add_paragraph()
run = para.add_run('2. 牛顿第一定律：')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = cell.add_paragraph()
run = para.add_run('    • 内容：一切物体...静止或匀速直线运动。')
set_chinese_font(run, 'SimSun', 12)

para = cell.add_paragraph()
run = para.add_run('    • 方法：实验 + 科学推理（理想实验法）。')
set_chinese_font(run, 'SimSun', 12)

para = cell.add_paragraph()
run = para.add_run('3. 惯性：')
set_chinese_font(run, 'SimHei', 12, bold=True)

para = cell.add_paragraph()
run = para.add_run('    • 定义：保持状态不变的性质。')
set_chinese_font(run, 'SimSun', 12)

para = cell.add_paragraph()
run = para.add_run('    • 决定因素：质量（与速度、受力无关）。')
set_chinese_font(run, 'SimSun', 12)

# 六、教学设计说明
add_heading_cn(doc, '六、教学设计说明', 1)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Inches(0.3)
run = para.add_run('本设计严格遵循"实验探究法"，利用斜面小车实验让学生经历从感性认识到理性推理的过程，有效解决了学生对"力与运动关系"的认知偏差。针对"惯性"这一难点，特意安排了概念辨析和生活实例应用环节，确保达成"能解决实际问题"的学习目标。')
set_chinese_font(run, 'SimSun', 12)

# Save document
doc.save(output_path)
print(f'Document generated successfully: {output_path}')